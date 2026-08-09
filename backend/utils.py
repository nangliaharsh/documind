import fitz  # PyMuPDF
import os
from typing import List
from docx import Document

def parse_pdf(file_path: str) -> List[str]:
    """Extract text from PDF page by page and chunk it."""
    doc = fitz.open(file_path)
    all_chunks = []

    for page_num, page in enumerate(doc):
        text = page.get_text()
        if text.strip():
            page_chunks = chunk_text(text, chunk_size=300, overlap=50)
            all_chunks.extend(page_chunks)

    doc.close()

    # Fallback: if still very few chunks, try full doc extraction
    if len(all_chunks) < 3:
        doc = fitz.open(file_path)
        full_text = ""
        for page in doc:
            full_text += page.get_text() + "\n"
        doc.close()
        all_chunks = chunk_text(full_text, chunk_size=300, overlap=50)

    print(f"Extracted {len(all_chunks)} chunks from {os.path.basename(file_path)}")
    return all_chunks if all_chunks else ["No readable text found in this PDF."]

def parse_docx(file_path: str) -> List[str]:
    """Extract text from DOCX file and chunk it."""
    doc = Document(file_path)
    full_text = ""

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            full_text += paragraph.text + "\n"

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    full_text += cell.text + "\n"

    chunks = chunk_text(full_text, chunk_size=300, overlap=50)
    print(f"Extracted {len(chunks)} chunks from {os.path.basename(file_path)}")
    return chunks if chunks else ["No readable text found in this DOCX."]

def parse_txt(file_path: str) -> List[str]:
    """Extract text from TXT file and chunk it."""
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        full_text = f.read()

    chunks = chunk_text(full_text, chunk_size=300, overlap=50)
    print(f"Extracted {len(chunks)} chunks from {os.path.basename(file_path)}")
    return chunks if chunks else ["No readable text found in this TXT."]

def parse_file(file_path: str) -> List[str]:
    """Auto-detect file type and parse accordingly."""
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return parse_pdf(file_path)
    elif ext == ".docx":
        return parse_docx(file_path)
    elif ext == ".txt":
        return parse_txt(file_path)
    else:
        print(f"Unsupported file type: {ext}")
        return [f"Unsupported file type: {ext}"]

def chunk_text(text: str, chunk_size: int = 300, overlap: int = 50) -> List[str]:
    """Split text into overlapping chunks by word count."""
    words = text.split()
    if not words:
        return []
    chunks = []
    for i in range(0, len(words), chunk_size - overlap):
        chunk = " ".join(words[i:i + chunk_size])
        if len(chunk.strip()) > 20:
            chunks.append(chunk)
    return chunks